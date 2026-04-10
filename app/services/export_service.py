from io import BytesIO
from textwrap import wrap

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.models.graph import PipelineResponse


def _to_pipeline_response(run_payload: dict) -> PipelineResponse:
    result_payload = run_payload.get("result", run_payload)
    return PipelineResponse(**result_payload)


def _iter_export_lines(result: PipelineResponse) -> list[str]:
    lines = [
        f"SynSoc AI Simulation Report",
        "",
        f"Topic: {result.topic}",
        f"Rounds: {result.simulation.total_rounds}",
        f"Turns: {result.simulation.total_turns}",
        f"Agents: {result.agents.total_agents}",
        "",
        "Executive Summary",
        result.report.executive_summary,
        "",
        "Key Findings",
    ]

    for finding in result.report.key_findings:
        lines.append(f"- {finding}")

    lines.extend(["", "Policy Recommendations"])
    for recommendation in result.report.policy_recommendations:
        lines.append(f"- {recommendation}")

    lines.extend(["", "Consensus Areas"])
    for area in result.report.consensus_areas:
        lines.append(f"- {area}")

    lines.extend(["", "Top Transcript Excerpts"])
    for turn in result.simulation.turns[:15]:
        lines.append(
            f"R{turn.round} {turn.agent_name} ({turn.represents}, {turn.stance}): {turn.message}"
        )

    return lines


def build_pdf_bytes(run_payload: dict) -> bytes:
    result = _to_pipeline_response(run_payload)
    lines = _iter_export_lines(result)

    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    x = 48
    y = height - 48

    pdf.setFont("Helvetica", 11)

    for line in lines:
        wrapped = wrap(line, width=100) if line else [""]
        for segment in wrapped:
            if y <= 48:
                pdf.showPage()
                pdf.setFont("Helvetica", 11)
                y = height - 48

            pdf.drawString(x, y, segment)
            y -= 14

    pdf.save()
    return buffer.getvalue()


def build_docx_bytes(run_payload: dict) -> bytes:
    result = _to_pipeline_response(run_payload)

    doc = Document()
    doc.add_heading("SynSoc AI Simulation Report", level=1)
    doc.add_paragraph(f"Topic: {result.topic}")
    doc.add_paragraph(f"Rounds: {result.simulation.total_rounds}")
    doc.add_paragraph(f"Turns: {result.simulation.total_turns}")
    doc.add_paragraph(f"Agents: {result.agents.total_agents}")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(result.report.executive_summary)

    doc.add_heading("Key Findings", level=2)
    for finding in result.report.key_findings:
        doc.add_paragraph(finding, style="List Bullet")

    doc.add_heading("Policy Recommendations", level=2)
    for recommendation in result.report.policy_recommendations:
        doc.add_paragraph(recommendation, style="List Bullet")

    doc.add_heading("Consensus Areas", level=2)
    for area in result.report.consensus_areas:
        doc.add_paragraph(area, style="List Bullet")

    doc.add_heading("Top Transcript Excerpts", level=2)
    for turn in result.simulation.turns[:15]:
        doc.add_paragraph(
            f"R{turn.round} {turn.agent_name} ({turn.represents}, {turn.stance}): {turn.message}"
        )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
